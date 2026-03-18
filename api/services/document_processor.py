from pathlib import Path

import openpyxl
import pdfplumber
from docx import Document as DocxDocument


class DocumentProcessor:
    """Pre-process documents before uploading to Dify.
    Especially important for Excel and scanned PDFs.
    """

    async def process(self, file_path: str) -> str:
        """Convert any document to clean markdown text."""
        ext = Path(file_path).suffix.lower()

        if ext in [".xlsx", ".xls"]:
            return self._process_excel(file_path)
        elif ext == ".pdf":
            return self._process_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            return self._process_docx(file_path)
        elif ext in [".txt", ".md"]:
            return Path(file_path).read_text(encoding="utf-8")
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    def _process_excel(self, file_path: str) -> str:
        """Convert Excel to Markdown tables."""
        wb = openpyxl.load_workbook(file_path, data_only=True)
        result = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            result.append(f"## Sheet: {sheet_name}\n")

            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                continue

            # Header row
            headers = [str(h) if h else "" for h in rows[0]]
            result.append("| " + " | ".join(headers) + " |")
            result.append("| " + " | ".join(["---"] * len(headers)) + " |")

            # Data rows
            for row in rows[1:]:
                cells = [str(c) if c is not None else "" for c in row]
                result.append("| " + " | ".join(cells) + " |")

            result.append("")

        return "\n".join(result)

    def _process_pdf(self, file_path: str) -> str:
        """Extract text from PDF, with OCR fallback for scanned docs."""
        text_parts = []

        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text and len(text.strip()) > 50:
                    text_parts.append(f"--- Trang {i + 1} ---\n{text}")
                else:
                    try:
                        import pytesseract

                        img = page.to_image(resolution=300)
                        ocr_text = pytesseract.image_to_string(
                            img.original, lang="vie+eng"
                        )
                        if ocr_text.strip():
                            text_parts.append(
                                f"--- Trang {i + 1} (OCR) ---\n{ocr_text}"
                            )
                    except Exception:
                        text_parts.append(f"--- Trang {i + 1} (không đọc được) ---")

        return "\n\n".join(text_parts)

    def _process_docx(self, file_path: str) -> str:
        """Extract text from DOCX preserving structure."""
        doc = DocxDocument(file_path)
        parts = []

        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                level = (
                    int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1
                )
                parts.append(f"{'#' * level} {para.text}")
            elif para.text.strip():
                parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append("| " + " | ".join(cells) + " |")
            if rows:
                header_sep = (
                    "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |"
                )
                rows.insert(1, header_sep)
                parts.append("\n".join(rows))

        return "\n\n".join(parts)


document_processor = DocumentProcessor()
